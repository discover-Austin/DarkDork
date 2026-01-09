#!/usr/bin/env python3
"""
DarkDork Advanced Export System
Support for PDF, DOCX, XML, and professional report generation
"""

import json
import csv
from datetime import datetime
from typing import List, Dict, Optional
import os


class ExportManager:
    """Manage all export formats"""

    def __init__(self):
        self.export_history = []

    def export_to_json(self, data: List[Dict], filename: str) -> bool:
        """Export data to JSON"""
        try:
            with open(filename, 'w') as f:
                json.dump({
                    'exported_at': datetime.now().isoformat(),
                    'tool': 'DarkDork',
                    'version': '1.0.0',
                    'record_count': len(data),
                    'data': data
                }, f, indent=2)

            self._record_export('json', filename, len(data))
            return True

        except Exception as e:
            print(f"Error exporting to JSON: {e}")
            return False

    def export_to_csv(self, data: List[Dict], filename: str) -> bool:
        """Export data to CSV"""
        try:
            if not data:
                return False

            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=data[0].keys())
                writer.writeheader()
                writer.writerows(data)

            self._record_export('csv', filename, len(data))
            return True

        except Exception as e:
            print(f"Error exporting to CSV: {e}")
            return False

    def export_to_xml(self, data: List[Dict], filename: str,
                     root_element: str = "darkdork_results") -> bool:
        """Export data to XML"""
        try:
            from xml.etree import ElementTree as ET
            from xml.dom import minidom

            root = ET.Element(root_element)
            root.set('exported_at', datetime.now().isoformat())
            root.set('tool', 'DarkDork')
            root.set('version', '1.0.0')
            root.set('record_count', str(len(data)))

            for item in data:
                record = ET.SubElement(root, 'record')
                for key, value in item.items():
                    element = ET.SubElement(record, key)
                    element.text = str(value) if value is not None else ''

            # Pretty print
            xml_string = minidom.parseString(ET.tostring(root)).toprettyxml(indent='  ')

            with open(filename, 'w', encoding='utf-8') as f:
                f.write(xml_string)

            self._record_export('xml', filename, len(data))
            return True

        except Exception as e:
            print(f"Error exporting to XML: {e}")
            return False

    def export_to_pdf(self, data: List[Dict], filename: str,
                     title: str = "DarkDork Report",
                     include_summary: bool = True) -> bool:
        """
        Export data to PDF
        Note: Requires reportlab package: pip install reportlab
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors

            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=30,
                alignment=1  # Center
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.3*inch))

            # Metadata
            meta_data = [
                ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
                ['Tool:', 'DarkDork Professional v1.0.0'],
                ['Total Records:', str(len(data))]
            ]

            meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
            meta_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ]))

            story.append(meta_table)
            story.append(Spacer(1, 0.5*inch))

            # Summary
            if include_summary and data:
                story.append(Paragraph("Summary", styles['Heading2']))
                story.append(Spacer(1, 0.2*inch))

                # Add summary statistics
                summary_text = f"This report contains {len(data)} records from DarkDork security assessment."
                story.append(Paragraph(summary_text, styles['Normal']))
                story.append(Spacer(1, 0.3*inch))

            # Data table
            if data:
                story.append(Paragraph("Detailed Results", styles['Heading2']))
                story.append(Spacer(1, 0.2*inch))

                # Create table data
                table_data = []

                # Headers
                if data:
                    headers = list(data[0].keys())
                    table_data.append(headers)

                    # Data rows (limit for PDF size)
                    for record in data[:100]:  # Limit to 100 records
                        row = [str(record.get(h, ''))[:50] for h in headers]  # Truncate long values
                        table_data.append(row)

                # Create table
                col_widths = [inch * (6.5 / len(headers)) for _ in headers]
                data_table = Table(table_data, colWidths=col_widths)

                data_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                ]))

                story.append(data_table)

                if len(data) > 100:
                    story.append(Spacer(1, 0.2*inch))
                    note = f"Note: Showing first 100 of {len(data)} records"
                    story.append(Paragraph(note, styles['Italic']))

            # Footer
            story.append(Spacer(1, 0.5*inch))
            footer_text = "Generated by DarkDork - Professional Google Dorking Tool"
            story.append(Paragraph(footer_text, styles['Italic']))

            # Build PDF
            doc.build(story)

            self._record_export('pdf', filename, len(data))
            return True

        except ImportError:
            print("Error: reportlab package not installed. Install with: pip install reportlab")
            return False
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return False

    def export_to_docx(self, data: List[Dict], filename: str,
                      title: str = "DarkDork Report") -> bool:
        """
        Export data to DOCX (Word)
        Note: Requires python-docx package: pip install python-docx
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Tool: DarkDork Professional v1.0.0")
            doc.add_paragraph(f"Total Records: {len(data)}")
            doc.add_paragraph('')  # Blank line

            # Summary section
            doc.add_heading('Summary', 1)
            summary = f"This report contains {len(data)} records from a DarkDork security assessment."
            doc.add_paragraph(summary)
            doc.add_paragraph('')

            # Detailed results
            doc.add_heading('Detailed Results', 1)

            if data:
                # Create table
                headers = list(data[0].keys())
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'

                # Header row
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    # Make header bold
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Data rows (limit for document size)
                for record in data[:500]:  # Limit to 500 records
                    row_cells = table.add_row().cells
                    for i, header in enumerate(headers):
                        value = record.get(header, '')
                        row_cells[i].text = str(value)[:100]  # Truncate long values

                if len(data) > 500:
                    doc.add_paragraph('')
                    note_para = doc.add_paragraph(f"Note: Showing first 500 of {len(data)} records")
                    note_para.italic = True

            # Footer
            doc.add_paragraph('')
            footer = doc.add_paragraph("Generated by DarkDork - Professional Google Dorking Tool")
            footer.italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Save document
            doc.save(filename)

            self._record_export('docx', filename, len(data))
            return True

        except ImportError:
            print("Error: python-docx package not installed. Install with: pip install python-docx")
            return False
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return False

    def export_to_html(self, data: List[Dict], filename: str,
                      title: str = "DarkDork Report",
                      include_css: bool = True) -> bool:
        """Export data to HTML"""
        try:
            html = []
            html.append('<!DOCTYPE html>')
            html.append('<html lang="en">')
            html.append('<head>')
            html.append('    <meta charset="UTF-8">')
            html.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
            html.append(f'    <title>{title}</title>')

            if include_css:
                html.append('    <style>')
                html.append('        body { font-family: Arial, sans-serif; margin: 40px; background: #f5f5f5; }')
                html.append('        .container { max-width: 1200px; margin: 0 auto; background: white; padding: 30px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }')
                html.append('        h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }')
                html.append('        .metadata { background: #ecf0f1; padding: 15px; margin: 20px 0; border-radius: 5px; }')
                html.append('        table { width: 100%; border-collapse: collapse; margin: 20px 0; }')
                html.append('        th { background: #3498db; color: white; padding: 12px; text-align: left; }')
                html.append('        td { padding: 10px; border-bottom: 1px solid #ddd; }')
                html.append('        tr:hover { background: #f5f5f5; }')
                html.append('        .footer { text-align: center; color: #7f8c8d; margin-top: 30px; font-style: italic; }')
                html.append('    </style>')

            html.append('</head>')
            html.append('<body>')
            html.append('    <div class="container">')
            html.append(f'        <h1>{title}</h1>')

            # Metadata
            html.append('        <div class="metadata">')
            html.append(f'            <p><strong>Generated:</strong> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>')
            html.append('            <p><strong>Tool:</strong> DarkDork Professional v1.0.0</p>')
            html.append(f'            <p><strong>Total Records:</strong> {len(data)}</p>')
            html.append('        </div>')

            # Data table
            if data:
                html.append('        <h2>Results</h2>')
                html.append('        <table>')
                html.append('            <thead>')
                html.append('                <tr>')

                headers = list(data[0].keys())
                for header in headers:
                    html.append(f'                    <th>{header}</th>')

                html.append('                </tr>')
                html.append('            </thead>')
                html.append('            <tbody>')

                for record in data:
                    html.append('                <tr>')
                    for header in headers:
                        value = record.get(header, '')
                        html.append(f'                    <td>{str(value)}</td>')
                    html.append('                </tr>')

                html.append('            </tbody>')
                html.append('        </table>')

            # Footer
            html.append('        <div class="footer">')
            html.append('            <p>Generated by DarkDork - Professional Google Dorking Tool</p>')
            html.append('        </div>')
            html.append('    </div>')
            html.append('</body>')
            html.append('</html>')

            with open(filename, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html))

            self._record_export('html', filename, len(data))
            return True

        except Exception as e:
            print(f"Error exporting to HTML: {e}")
            return False

    def export_to_markdown(self, data: List[Dict], filename: str,
                          title: str = "DarkDork Report") -> bool:
        """Export data to Markdown"""
        try:
            md = []
            md.append(f"# {title}\n")
            md.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ")
            md.append(f"**Tool:** DarkDork Professional v1.0.0  ")
            md.append(f"**Total Records:** {len(data)}\n")

            if data:
                md.append("## Results\n")

                # Create markdown table
                headers = list(data[0].keys())

                # Headers
                md.append('| ' + ' | '.join(headers) + ' |')
                md.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')

                # Data rows
                for record in data:
                    row = []
                    for header in headers:
                        value = str(record.get(header, ''))
                        # Escape pipe characters
                        value = value.replace('|', '\\|')
                        row.append(value[:100])  # Truncate long values
                    md.append('| ' + ' | '.join(row) + ' |')

            md.append("\n---")
            md.append("\n*Generated by DarkDork - Professional Google Dorking Tool*")

            with open(filename, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md))

            self._record_export('markdown', filename, len(data))
            return True

        except Exception as e:
            print(f"Error exporting to Markdown: {e}")
            return False

    def _record_export(self, format_type: str, filename: str, record_count: int):
        """Record export in history"""
        self.export_history.append({
            'timestamp': datetime.now().isoformat(),
            'format': format_type,
            'filename': filename,
            'record_count': record_count
        })

    def get_export_history(self) -> List[Dict]:
        """Get export history"""
        return self.export_history


# Example usage
def example_exports():
    """Example export usage"""

    print("DarkDork Advanced Export System")
    print("="*50)

    # Sample data
    sample_data = [
        {
            'timestamp': '2026-01-09 10:30:00',
            'dork': 'filetype:pdf confidential',
            'target': 'example.com',
            'status': 'executed',
            'severity': 'high'
        },
        {
            'timestamp': '2026-01-09 10:32:00',
            'dork': 'inurl:admin intitle:login',
            'target': 'example.com',
            'status': 'executed',
            'severity': 'medium'
        },
        {
            'timestamp': '2026-01-09 10:35:00',
            'dork': 'intitle:"index of" .env',
            'target': 'example.com',
            'status': 'executed',
            'severity': 'critical'
        }
    ]

    exporter = ExportManager()

    # Export to different formats
    print("\nExporting to multiple formats...")

    formats = [
        ('JSON', 'example_export.json', lambda: exporter.export_to_json(sample_data, 'example_export.json')),
        ('CSV', 'example_export.csv', lambda: exporter.export_to_csv(sample_data, 'example_export.csv')),
        ('XML', 'example_export.xml', lambda: exporter.export_to_xml(sample_data, 'example_export.xml')),
        ('HTML', 'example_export.html', lambda: exporter.export_to_html(sample_data, 'example_export.html')),
        ('Markdown', 'example_export.md', lambda: exporter.export_to_markdown(sample_data, 'example_export.md')),
        ('PDF', 'example_export.pdf', lambda: exporter.export_to_pdf(sample_data, 'example_export.pdf')),
        ('DOCX', 'example_export.docx', lambda: exporter.export_to_docx(sample_data, 'example_export.docx')),
    ]

    for format_name, filename, export_func in formats:
        success = export_func()
        status = "✓ Success" if success else "✗ Failed"
        print(f"  {format_name:10} → {filename:25} {status}")

    print("\n✓ Export examples complete!")


if __name__ == '__main__':
    example_exports()
